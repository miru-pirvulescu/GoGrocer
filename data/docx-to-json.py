from docx import Document
import json

def read(fileName, index):
    doc = Document(fileName)
    table = doc.tables[index]
    elements = []
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        
        if i == 0:
            keys = tuple(text)
            continue

        raw_data = list(dict(zip(keys, text)).values())[:-1]
        elements.append(raw_data)
    return elements

    # return data

def main():
    from docx import Document
import json

def read(fileName, index):
    doc = Document(fileName)
    table = doc.tables[index]
    elements = []
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        
        if i == 0:
            keys = tuple(text)
            continue

        raw_data = list(dict(zip(keys, text)).values())[:-1]
        elements.append(raw_data)
    return elements



def toJSON(preDoc, postDoc):
    doc = Document(preDoc)
    bar = "barman"
    data = {"author" : "Silvia N. Mirica",
            "bookName" : "recipes",
            "contents": {
                "id" : 1,
                "recipeName": "foo",
                "ingredients": bar
            }
        }

    recipes = open (postDoc, 'a')
    
    recipes.write(json.dumps(data, indent=4))

def main():
    toJSON('recipes.docx', 'recipes.json')

if __name__ == "__main__":
    main()
